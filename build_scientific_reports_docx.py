from __future__ import annotations

from collections import OrderedDict
from pathlib import Path
import re
import shutil
import subprocess
from urllib.parse import quote
from urllib.request import Request, urlopen

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
import pymupdf
import pypandoc


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "final_review_v1.tex"
OUTPUT = ROOT / "final_review_v1_scientific_reports.docx"
NOTES = ROOT / "final_review_v1_scientific_reports_conversion_notes.txt"
BUILD = ROOT / "_docx_build"
ASSETS = BUILD / "public_figures"
FIGURES = BUILD / "figures"
PREPARED_TEX = BUILD / "final_review_v1_prepared.tex"
RAW_DOCX = BUILD / "final_review_v1_raw.docx"

REPOSITORY_RAW = (
    "https://raw.githubusercontent.com/"
    "DekuLabUO/Neural-Probe-Strain/main/"
)

PUBLIC_ASSETS = {
    "Figure1.pdf": "Images/Figure1_MEA_Geometry.pdf",
    "Figure2.png": "Images/Figure2_Regression_Validation.png",
    "Figure3.pdf": "Images/Figure3_LSMean_Strain.pdf",
    "Figure4.pdf": "Images/Figure4_TwoWay_Interaction.pdf",
    "Figure5.pdf": "Images/Figure5_ThreeWay_Interaction.pdf",
    "Figure6.pdf": "Images/Figure6_FourWay_Interaction.pdf",
    "Figure7a.pdf": (
        "Images/Figure8_ShortWide_Device/"
        "3D_Volumetric_Strain_Distribution.pdf"
    ),
    "Figure7b.pdf": (
        "Images/Figure8_ShortWide_Device/Equivalent_Strain_Profiles.pdf"
    ),
    "Figure7d.pdf": (
        "Images/Figure8_ShortWide_Device/Heatmap_Average_Strain.pdf"
    ),
    "Figure7e.pdf": (
        "Images/Figure8_ShortWide_Device/Quantitative_Comparison.pdf"
    ),
    "Figure8a.pdf": (
        "Images/Figure7_LongThin_Device/"
        "3D_Volumetric_Strain_Distribution.pdf"
    ),
    "Figure8b.pdf": (
        "Images/Figure7_LongThin_Device/Equivalent_Strain_Profiles.pdf"
    ),
    "Figure8d.pdf": (
        "Images/Figure7_LongThin_Device/Heatmap_Average_Strain.pdf"
    ),
    "Figure8e.pdf": (
        "Images/Figure7_LongThin_Device/Quantitative_Comparison.pdf"
    ),
}

FIGURE_PATHS = {
    "figures/Figure 1.png": "figures/Figure1.png",
    "figures/Figure2_v2.png": "figures/Figure2.png",
    "figures/Figure3_Not_Bold.pdf": "figures/Figure3.png",
    "figures/Figure_4.pdf": "figures/Figure4.png",
    "figures/Figure_5.pdf": "figures/Figure5.png",
    "figures/Figure_6.pdf": "figures/Figure6.png",
    "figures/Fig7.pdf": "figures/Figure7.jpg",
    "figures/Figur-8.pdf": "figures/Figure8.jpg",
}

REFERENCE_NUMBERS = {
    "fig:fig1": "1",
    "fig:fig2": "2",
    "fig:fig3": "3",
    "fig:fig4": "4",
    "fig:fig5": "5",
    "fig:fig6": "6",
    "fig:fig7": "7",
    "fig:fig8": "8",
    "Table 1": "1",
    "tab:anova": "2",
    "Table 2": "3",
}


def download_asset(name: str, repository_path: str) -> Path:
    destination = ASSETS / name
    if destination.exists() and destination.stat().st_size:
        return destination

    destination.parent.mkdir(parents=True, exist_ok=True)
    url = REPOSITORY_RAW + quote(repository_path, safe="/")
    request = Request(url, headers={"User-Agent": "Scientific-Reports-DOCX"})
    with urlopen(request) as response, destination.open("wb") as output:
        shutil.copyfileobj(response, output)
    return destination


def pdf_image(path: Path, max_width: int = 1800) -> Image.Image:
    document = pymupdf.open(path)
    page = document[0]
    scale = max(1.0, min(2.4, max_width / page.rect.width))
    pixmap = page.get_pixmap(
        matrix=pymupdf.Matrix(scale, scale),
        alpha=False,
    )
    image = Image.frombytes("RGB", (pixmap.width, pixmap.height), pixmap.samples)
    if image.width > max_width:
        height = round(image.height * max_width / image.width)
        image = image.resize((max_width, height), Image.Resampling.LANCZOS)
    document.close()
    return image


def render_pdf(source: Path, destination: Path, max_width: int = 1800) -> None:
    image = pdf_image(source, max_width=max_width)
    image.save(destination, "PNG", optimize=True, compress_level=9)


def fit_image(image: Image.Image, width: int, height: int) -> Image.Image:
    fitted = image.copy()
    fitted.thumbnail((width, height), Image.Resampling.LANCZOS)
    return fitted


def panel_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    filename = "arialbd.ttf" if bold else "arial.ttf"
    path = Path("C:/Windows/Fonts") / filename
    if path.exists():
        return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def compose_figure(
    destination: Path,
    panels: dict[str, Path | None],
) -> None:
    canvas_width = 1800
    canvas_height = 2060
    margin = 45
    gap = 30
    cell_width = (canvas_width - 2 * margin - gap) // 2
    cell_height = 620
    canvas = Image.new("RGB", (canvas_width, canvas_height), "white")
    draw = ImageDraw.Draw(canvas)
    label_font = panel_font(34, bold=True)
    note_font = panel_font(25)

    locations = {
        "a": (margin, margin),
        "b": (margin + cell_width + gap, margin),
        "c": (margin, margin + cell_height + gap),
        "d": (margin + cell_width + gap, margin + cell_height + gap),
        "e": ((canvas_width - cell_width) // 2, margin + 2 * (cell_height + gap)),
    }

    for label, (left, top) in locations.items():
        draw.text((left, top), f"({label})", fill="black", font=label_font)
        content_top = top + 50
        source = panels[label]
        if source is None:
            box = (
                left + 40,
                content_top + 50,
                left + cell_width - 40,
                top + cell_height - 40,
            )
            draw.rectangle(box, outline=(140, 140, 140), width=3)
            lines = [
                "Simulation geometry panel",
                "not present in the supplied TeX",
                "or public figure repository.",
            ]
            line_height = 38
            start_y = (box[1] + box[3] - line_height * len(lines)) // 2
            for index, line in enumerate(lines):
                bounds = draw.textbbox((0, 0), line, font=note_font)
                text_width = bounds[2] - bounds[0]
                draw.text(
                    ((box[0] + box[2] - text_width) // 2, start_y + index * line_height),
                    line,
                    fill=(80, 80, 80),
                    font=note_font,
                )
            continue

        image = fit_image(pdf_image(source, max_width=1500), cell_width, cell_height - 65)
        image_left = left + (cell_width - image.width) // 2
        image_top = content_top + (cell_height - 65 - image.height) // 2
        canvas.paste(image, (image_left, image_top))

    canvas.save(destination, "JPEG", quality=91, optimize=True, progressive=True)


def prepare_figures() -> None:
    ASSETS.mkdir(parents=True, exist_ok=True)
    FIGURES.mkdir(parents=True, exist_ok=True)
    downloaded = {
        name: download_asset(name, repository_path)
        for name, repository_path in PUBLIC_ASSETS.items()
    }

    render_pdf(downloaded["Figure1.pdf"], FIGURES / "Figure1.png")
    figure2 = Image.open(downloaded["Figure2.png"]).convert("RGB")
    figure2.save(FIGURES / "Figure2.png", "PNG", optimize=True)
    for number in range(3, 7):
        render_pdf(downloaded[f"Figure{number}.pdf"], FIGURES / f"Figure{number}.png")

    # The public folder names for these datasets conflict with their plotted
    # regional trends. Match panels to the manuscript captions by the data:
    # Figure 7 has tip-dominant strain; Figure 8 has top-dominant strain.
    compose_figure(
        FIGURES / "Figure7.jpg",
        {
            "a": downloaded["Figure8a.pdf"],
            "b": downloaded["Figure8b.pdf"],
            "c": None,
            "d": downloaded["Figure8d.pdf"],
            "e": downloaded["Figure8e.pdf"],
        },
    )
    compose_figure(
        FIGURES / "Figure8.jpg",
        {
            "a": downloaded["Figure7a.pdf"],
            "b": downloaded["Figure7b.pdf"],
            "c": None,
            "d": downloaded["Figure7d.pdf"],
            "e": downloaded["Figure7e.pdf"],
        },
    )


def citation_order(text: str) -> OrderedDict[str, int]:
    order: OrderedDict[str, int] = OrderedDict()
    for match in re.finditer(r"\\cite\{([^}]+)\}", text):
        for key in match.group(1).split(","):
            key = key.strip()
            if key and key not in order:
                order[key] = len(order) + 1
    return order


def compact_numbers(numbers: list[int]) -> str:
    if not numbers:
        return ""
    numbers = sorted(set(numbers))
    groups: list[str] = []
    start = previous = numbers[0]
    for number in numbers[1:] + [numbers[-1] + 2]:
        if number == previous + 1:
            previous = number
            continue
        if start == previous:
            groups.append(str(start))
        elif previous == start + 1:
            groups.extend((str(start), str(previous)))
        else:
            groups.append(f"{start}\u2013{previous}")
        start = previous = number
    return ", ".join(groups)


def replace_citations(text: str, order: OrderedDict[str, int]) -> str:
    def replacement(match: re.Match[str]) -> str:
        keys = [key.strip() for key in match.group(1).split(",")]
        numbers = [order[key] for key in keys]
        return f"[{compact_numbers(numbers)}]"

    text = re.sub(r"\\cite\{([^}]+)\}", replacement, text)
    return re.sub(r"(\])(?=[A-Za-z])", r"\1 ", text)


def replace_cross_references(text: str) -> str:
    def replacement(match: re.Match[str]) -> str:
        label = match.group(1)
        return REFERENCE_NUMBERS.get(label, label)

    return re.sub(r"\\ref\{([^}]+)\}", replacement, text)


def build_reference_fallback(order: OrderedDict[str, int]) -> str:
    items = "\n".join(
        f"\\item \\texttt{{{key}}}" for key in order
    )
    return (
        "\\section*{References}\n"
        "\\noindent\\textbf{[REFERENCE METADATA REQUIRED: The source calls "
        "\\texttt{reffinal.bib}, but that file was not supplied. The numbered "
        "entries below preserve every original BibTeX key and its citation order.]}\n"
        "\\begin{enumerate}\n"
        f"{items}\n"
        "\\end{enumerate}\n"
    )


def prepare_tex() -> OrderedDict[str, int]:
    text = SOURCE.read_text(encoding="utf-8")
    order = citation_order(text)

    abstract_match = re.search(
        r"\\begin\{abstract\}(.*?)\\end\{abstract\}",
        text,
        flags=re.DOTALL,
    )
    if not abstract_match:
        raise RuntimeError("Abstract not found in source manuscript")
    abstract = abstract_match.group(1).strip()
    abstract = re.sub(
        r"\\textbf\{Keywords:\}.*$",
        "",
        abstract,
        flags=re.DOTALL,
    ).strip()

    keywords_match = re.search(r"\\keywords\{([^}]*)\}", text)
    if not keywords_match:
        raise RuntimeError("Keywords not found in source manuscript")
    keywords = keywords_match.group(1).strip()

    text = text[: abstract_match.start()] + text[abstract_match.end() :]
    text = re.sub(r"^\\keywords\{[^}]*\}\s*", "", text, flags=re.MULTILINE)
    text = re.sub(
        r"^\\author(?:\[[^]]*\])?\{[^}]*\}\s*",
        "",
        text,
        flags=re.MULTILINE,
    )
    text = re.sub(
        r"^\\affil(?:\[[^]]*\])?\{[^}]*\}\s*",
        "",
        text,
        flags=re.MULTILINE,
    )

    authors = (
        "\\author{Rubiya Yasmin\\textsuperscript{1}}\n"
        "\\author{Yupeng Wu\\textsuperscript{1}}\n"
        "\\author{Benjamin Aleman\\textsuperscript{2,3,4}}\n"
        "\\author{Felix Deku\\textsuperscript{1,*}}\n"
    )
    title_match = re.search(r"\\title\{[^}]*\}", text)
    if not title_match:
        raise RuntimeError("Title not found in source manuscript")
    text = text[: title_match.end()] + "\n" + authors + text[title_match.end() :]

    title_page = (
        "\n\\begin{center}\n"
        "\\textsuperscript{1}Department of Bioengineering, University of "
        "Oregon, Eugene, OR, United States of America\\\\\n"
        "\\textsuperscript{2}Department of Physics, University of Oregon, "
        "Eugene, OR 97403, USA\\\\\n"
        "\\textsuperscript{3}Materials Science Institute, University of "
        "Oregon, Eugene, OR 97403, USA\\\\\n"
        "\\textsuperscript{4}Center for Optical, Molecular, and Quantum "
        "Science, University of Oregon, Eugene, OR 97403, USA\\\\\n"
        "\\textsuperscript{*}Corresponding author: fdeku@uoregon.edu\n"
        "\\end{center}\n"
        "\\section*{Abstract}\n"
        f"{abstract}\n\n"
        f"\\noindent\\textbf{{Keywords:}} {keywords}\n"
        "\\newpage\n"
    )
    text = text.replace("\\maketitle", "\\maketitle" + title_page, 1)

    for original, prepared in FIGURE_PATHS.items():
        text = text.replace(original, prepared)

    text = replace_citations(text, order)
    text = replace_cross_references(text)
    text = text.replace(
        r"(1~\mathrm{mm}\times1~\mathrm{mm}\times10~\mathrm{mm})",
        r"($1~\mathrm{mm}\times1~\mathrm{mm}\times10~\mathrm{mm}$)",
    )
    text = re.sub(
        r"\\tag\{(\d+)\}",
        lambda match: rf"\qquad\text{{({match.group(1)})}}",
        text,
    )
    text = re.sub(r"Equation~(\d+)", r"Equation~(\1)", text)
    text = text.replace("\\section{Result Analysis}", "\\section{Results}")

    data_match = re.search(
        r"\\section\*\{Data Availability\s*\}(.*?)"
        r"(?=\\section\*\{Funding\})",
        text,
        flags=re.DOTALL,
    )
    funding_match = re.search(
        r"\\section\*\{Funding\}(.*?)"
        r"(?=%?\\bibliographystyle|\\bibliography\{reffinal\})",
        text,
        flags=re.DOTALL,
    )
    if not data_match or not funding_match:
        raise RuntimeError("Data Availability or Funding section not found")

    body = text[: data_match.start()].rstrip()
    data_text = data_match.group(1).strip()
    funding_text = funding_match.group(1).strip()
    tail = (
        "\n\n\\section*{Data Availability}\n"
        f"{data_text}\n\n"
        f"{build_reference_fallback(order)}\n"
        "\\section*{Funding}\n"
        f"{funding_text}\n\n"
        "\\section*{Author Contributions}\n"
        "\\textbf{[AUTHOR INPUT REQUIRED: Describe each author's "
        "contributions to the work.]}\n\n"
        "\\section*{Competing Interests}\n"
        "\\textbf{[AUTHOR INPUT REQUIRED: Provide an explicit competing "
        "interests declaration for each author.]}\n\n"
        "\\end{document}\n"
    )
    PREPARED_TEX.write_text(body + tail, encoding="utf-8")
    return order


def set_style_font(style, name: str, size: float, bold: bool | None = None) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        style.font.bold = bold
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.add_run("Page ")
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, end))


def enable_line_numbers(section) -> None:
    section_properties = section._sectPr
    existing = section_properties.find(qn("w:lnNumType"))
    if existing is not None:
        section_properties.remove(existing)
    line_numbers = OxmlElement("w:lnNumType")
    line_numbers.set(qn("w:countBy"), "1")
    line_numbers.set(qn("w:distance"), "360")
    line_numbers.set(qn("w:restart"), "continuous")
    section_properties.append(line_numbers)


def prefix_caption(paragraph, prefix: str) -> None:
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    bold = OxmlElement("w:b")
    properties.append(bold)
    run.append(properties)
    text = OxmlElement("w:t")
    text.set(qn("xml:space"), "preserve")
    text.text = prefix
    run.append(text)

    children = list(paragraph._p)
    insertion_index = 0
    for index, child in enumerate(children):
        if child.tag != qn("w:r"):
            continue
        visible_text = "".join(child.itertext()).strip()
        if visible_text:
            insertion_index = index
            break
        insertion_index = index + 1
    paragraph._p.insert(insertion_index, run)


def repeat_table_header(row) -> None:
    row_properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    row_properties.append(header)


def postprocess_docx() -> None:
    document = Document(RAW_DOCX)
    document.core_properties.title = (
        "Microelectrode Geometry Governs Micromotion-induced Tissue Strain "
        "at the Brain-Implant Interface"
    )
    document.core_properties.subject = "Scientific Reports manuscript"
    document.core_properties.author = (
        "Rubiya Yasmin; Yupeng Wu; Benjamin Aleman; Felix Deku"
    )

    body_style_names = ("Normal", "Body Text", "First Paragraph")
    for style_name in body_style_names:
        if style_name in document.styles:
            style = document.styles[style_name]
            set_style_font(style, "Times New Roman", 12)
            style.paragraph_format.line_spacing = 2
            style.paragraph_format.space_after = Pt(0)

    style_settings = {
        "Title": (16, True, 1.15),
        "Author": (11, False, 1.0),
        "Heading 1": (14, True, 1.15),
        "Heading 2": (12, True, 1.15),
        "Heading 3": (12, False, 1.15),
        "Image Caption": (10, False, 1.0),
        "Table Caption": (10, False, 1.0),
    }
    for style_name, (size, bold, spacing) in style_settings.items():
        if style_name not in document.styles:
            continue
        style = document.styles[style_name]
        set_style_font(style, "Times New Roman", size, bold=bold)
        style.paragraph_format.line_spacing = spacing
        style.paragraph_format.keep_with_next = style_name.startswith("Heading")
        style.paragraph_format.space_before = Pt(8 if "Heading" in style_name else 3)
        style.paragraph_format.space_after = Pt(3)
    if "Heading 3" in document.styles:
        document.styles["Heading 3"].font.italic = True

    if "Reference Placeholder" not in document.styles:
        reference_style = document.styles.add_style(
            "Reference Placeholder",
            WD_STYLE_TYPE.PARAGRAPH,
        )
        set_style_font(reference_style, "Times New Roman", 10)
        reference_style.paragraph_format.line_spacing = 1
        reference_style.paragraph_format.first_line_indent = Inches(-0.25)
        reference_style.paragraph_format.left_indent = Inches(0.25)

    for section in document.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        enable_line_numbers(section)
        footer = section.footer
        footer.is_linked_to_previous = False
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.clear()
        add_page_number(footer_paragraph)

    figure_number = 0
    table_number = 0
    in_reference_list = False
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        style_name = paragraph.style.name

        if text == "Introduction":
            paragraph.paragraph_format.page_break_before = True
        if text == "References":
            in_reference_list = True
        elif in_reference_list and text in {
            "Funding",
            "Author Contributions",
            "Competing Interests",
        }:
            in_reference_list = False

        if style_name == "Image Caption":
            figure_number += 1
            prefix_caption(paragraph, f"Figure {figure_number}. ")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif style_name == "Table Caption":
            table_number += 1
            prefix_caption(paragraph, f"Table {table_number}. ")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.keep_with_next = True

        if paragraph._p.xpath(".//w:drawing"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.line_spacing = 1

        if style_name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True

        if in_reference_list and re.fullmatch(r"\d+\.\s*\w+", text):
            paragraph.style = document.styles["Reference Placeholder"]

        if "[AUTHOR INPUT REQUIRED:" in text or "[REFERENCE METADATA REQUIRED:" in text:
            for run in paragraph.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                run.font.bold = True

    for table in document.tables:
        if "Table Grid" in document.styles:
            table.style = "Table Grid"
        table.autofit = True
        if table.rows:
            repeat_table_header(table.rows[0])
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1
                    paragraph.paragraph_format.space_after = Pt(0)
                    if row_index < len(table.rows) - 1:
                        paragraph.paragraph_format.keep_with_next = True
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(8.5)
                        run.font.bold = row_index == 0

    document.save(OUTPUT)


def strip_tex_for_count(text: str) -> str:
    text = re.sub(r"%[^\n]*", " ", text)
    text = re.sub(
        r"\\begin\{(?:figure|table)\}.*?\\end\{(?:figure|table)\}",
        " ",
        text,
        flags=re.DOTALL,
    )
    text = re.sub(r"\\cite\{[^}]*\}", " ", text)
    text = re.sub(r"\\[A-Za-z*]+(?:\[[^]]*\])?", " ", text)
    text = re.sub(r"[^A-Za-z0-9'-]+", " ", text)
    return text


def word_count(text: str) -> int:
    return len(strip_tex_for_count(text).split())


def write_notes(order: OrderedDict[str, int]) -> None:
    source = SOURCE.read_text(encoding="utf-8")
    title = re.search(r"\\title\{([^}]*)\}", source).group(1)
    abstract = re.search(
        r"\\begin\{abstract\}(.*?)\\end\{abstract\}",
        source,
        flags=re.DOTALL,
    ).group(1)
    abstract = re.sub(r"\\textbf\{Keywords:\}.*$", "", abstract, flags=re.DOTALL)

    introduction = re.search(
        r"\\section\{Introduction\}(.*?)(?=\\section\{Methods\})",
        source,
        flags=re.DOTALL,
    ).group(1)
    results = re.search(
        r"\\section\{Result Analysis\}(.*?)(?=\\section\{Discussion\})",
        source,
        flags=re.DOTALL,
    ).group(1)
    discussion = re.search(
        r"\\section\{Discussion\}(.*?)(?=\\section\{Conclusion\})",
        source,
        flags=re.DOTALL,
    ).group(1)
    conclusion = re.search(
        r"\\section\{Conclusion\}(.*?)(?=\\section\*\{Data Availability)",
        source,
        flags=re.DOTALL,
    ).group(1)
    main_count = sum(
        word_count(part)
        for part in (introduction, results, discussion, conclusion)
    )

    notes = f"""Scientific Reports DOCX conversion audit
===========================================

Output
------
DOCX: {OUTPUT}
Source TeX (unchanged): {SOURCE}

Preserved
---------
- Native Word headings, editable tables, and editable Word equations.
- Page numbers and continuous line numbers for review.
- All {len(order)} unique citation keys, mapped to numerical citations in first-use order.
- Eight inline figures and their original legends.

Source dependencies not supplied
--------------------------------
- reffinal.bib is absent. The References section therefore contains numbered BibTeX-key placeholders. Replace these with full Nature-style entries before submission.
- The eight figure files named by the TeX are absent. Figures 1-6 were recovered from the authors' public data repository.
- Figures 7-8 were reconstructed from the matching public data panels. Panel (c), described in each source legend, was not present in the supplied files or public repository and is explicitly marked as unavailable.

Required author input
---------------------
- Complete the highlighted Author Contributions statement.
- Complete the highlighted Competing Interests statement.

Scientific Reports guideline checks (26 August 2026)
----------------------------------------------------
- Title: {word_count(title)} words (recommended maximum: 20).
- Abstract: {word_count(abstract)} words (recommended maximum: 200).
- Main text excluding Methods, references, tables, and figure legends: approximately {main_count} words (recommended maximum: 4,500).
- Keywords: 6 (maximum allowed: 6).
- Display items: 8 figures + 3 tables = 11 (journal limit stated as 8).
- Unique references: {len(order)} (journal recommends about 60, though not strictly enforced).

No scientific prose, numerical results, or source tables were silently shortened to meet these limits.
"""
    NOTES.write_text(notes, encoding="utf-8")


def convert() -> None:
    BUILD.mkdir(parents=True, exist_ok=True)
    prepare_figures()
    order = prepare_tex()

    pandoc = pypandoc.get_pandoc_path()
    command = [
        pandoc,
        str(PREPARED_TEX),
        "--from=latex",
        "--to=docx",
        "--standalone",
        f"--resource-path={BUILD}",
        f"--output={RAW_DOCX}",
    ]
    subprocess.run(command, check=True)
    postprocess_docx()
    write_notes(order)


def main() -> None:
    convert()
    print(OUTPUT)
    print(NOTES)


if __name__ == "__main__":
    main()