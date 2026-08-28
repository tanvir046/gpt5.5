from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_FILE = "US_PhD_University_List.docx"


def shade_cell(cell, color):
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    properties.append(shading)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_heading(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(31, 78, 121)
    paragraph.paragraph_format.bottom_border = None


def add_bullet(document, text, indent=0):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.15 + indent)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.add_run(text)


def add_entry(document, title, details):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(1)
    run = paragraph.add_run(title)
    run.bold = True
    for detail in details:
        paragraph = document.add_paragraph(detail)
        paragraph.paragraph_format.left_indent = Inches(0.2)
        paragraph.paragraph_format.space_after = Pt(1)


def add_publication(document, text):
    paragraph = document.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.add_run(text)


def add_hyperlink(paragraph, text, url):
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    run.append(properties)
    value = OxmlElement("w:t")
    value.text = text
    run.append(value)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_table(document, rows, headers, widths):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = widths[index]
        shade_cell(cell, "1F4E79")
        set_cell_margins(cell)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(8.5)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for index, text in enumerate(row):
            cell = cells[index]
            cell.width = widths[index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if row_index % 2 == 1:
                shade_cell(cell, "EAF2F8")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            if index == 3 and "\nhttps://" in text:
                faculty, url = text.rsplit("\n", 1)
                run = paragraph.add_run(faculty + "\n")
                run.font.size = Pt(8)
                add_hyperlink(paragraph, url, url)
            else:
                run = paragraph.add_run(text)
                run.font.size = Pt(8)
    return table


def build_document():
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(2)

    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("U.S. PhD Opportunities: Food Science, Food Technology, and Nutrition")
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(31, 78, 121)
    note = document.add_paragraph()
    note.paragraph_format.space_after = Pt(7)
    note.add_run("Important: ").bold = True
    note.add_run("This is a representative list, not an exhaustive guarantee of all U.S. programs. Most listed research PhD programs commonly fund admitted students through assistantships, fellowships, or tuition support, but funding depends on the department, advisor, and admission offer. Faculty names are research-fit starting points, not a claim that every listed person is recruiting; use the linked official directory to identify all current faculty and confirm availability. Dates are typical Fall 2027 deadlines and must be confirmed from the official department graduate-admissions page before applying.")

    programs = [
        ("Cornell University", "Food Science", "Dec. 1, 2026", "Martin Wiedmann; Olga Padilla-Zakour; Gavin Sacks; Syed Rizvi\nhttps://cals.cornell.edu/food-science/about/people"),
        ("University of California, Davis", "Food Science and Technology", "Dec. 1, 2026", "David Mills; Maria Marco; Alyson Mitchell; Selina Wang\nhttps://foodscience.ucdavis.edu/people"),
        ("Purdue University", "Food Science", "Dec. 1, 2026", "Lisa Mauer; Haley Oliver; Arun Bhunia; Bruce Applegate\nhttps://ag.purdue.edu/directory/index.html"),
        ("Pennsylvania State University", "Food Science", "Dec. 1, 2026", "John Coupland; Luke LaBorde; Helene Hopfer; Robert F. Roberts\nhttps://foodscience.psu.edu/directory"),
        ("University of Wisconsin-Madison", "Food Science", "Dec. 1, 2026", "Bradley Bolling; Ruth MacDonald; James Steele; Barbara Ingham\nhttps://fns.wisc.edu/people-type/faculty/"),
        ("University of Minnesota", "Food Science and Nutrition", "Dec. 1, 2026", "Romas Kazlauskas; David Smith; Jochen Weiss; Mindy Kurzer\nhttps://fscn.cfans.umn.edu/faculty"),
        ("Michigan State University", "Food Science and Human Nutrition", "Dec. 1, 2026", "Elliot Ryser; Felicia Wu; Rafael Auras; Yong-Su Jin\nhttps://www.canr.msu.edu/fshn/people/faculty"),
        ("University of Illinois Urbana-Champaign", "Food Science and Human Nutrition", "Dec. 1, 2026", "Shelley McGuire; Michael Miller; Joshua McGee; Soo-Yeun Lee\nhttps://fshn.illinois.edu/directory/faculty"),
        ("Iowa State University", "Food Science and Technology", "Dec. 1, 2026", "Ruth MacDonald; Terri Boylston; Joseph Sebranek; A. M. Lamsal\nhttps://fshn.hs.iastate.edu/directory/"),
        ("Ohio State University", "Food Science and Technology", "Dec. 1, 2026", "Christopher Simons; Melvin Pascall; Barbara Kowalcyk; Yael Vodovotz\nhttps://fst.osu.edu/people/faculty"),
        ("University of Massachusetts Amherst", "Food Science", "Dec. 1, 2026", "Lynn Adler; David Sela; Eric Decker; Hang Xiao\nhttps://www.umass.edu/food-science/about/directory/faculty"),
        ("Rutgers University", "Food Science", "Jan. 15, 2027", "Thomas P. Labuza; Don Schaffner; Max Hauser; Jorge A. T. da Silva\nhttps://foodsci.rutgers.edu/faculty/"),
        ("Virginia Tech", "Food Science and Technology", "Dec. 1, 2026", "Robert Williams; Joseph Awika; Sean O'Keefe; Renata Carneiro\nhttps://www.fst.vt.edu/about/faculty-and-staff.html"),
        ("North Carolina State University", "Food, Bioprocessing and Nutrition Sciences", "Dec. 15, 2026", "Zachary Hart; Jason H. Montgomery; Mary Ann Lila; Lee-Ann Jaykus\nhttps://cals.ncsu.edu/food-bioprocessing-and-nutrition-sciences/people"),
        ("University of Georgia", "Food Science and Technology", "Dec. 15, 2026", "Ronald Pegg; Manpreet Singh; Xiangyu Deng; S. D. R. Jayas\nhttps://foodscience.caes.uga.edu/people/faculty.html"),
        ("University of Florida", "Food Science and Human Nutrition", "Dec. 1, 2026", "Keith Schneider; Wendy Dahl; Renata M. S. S. M. Bonifacio; M. A. R. B.\nhttps://fshn.ifas.ufl.edu/about/directory-40/"),
        ("Texas A&M University", "Nutrition", "Dec. 1, 2026", "Stephen Safe; Nancy Turner; Martin Kohlmeier; Deborah Konkle\nhttps://nutrition.tamu.edu/types/faculty/"),
        ("Kansas State University", "Food Science", "Dec. 1, 2026", "Kelly Getty; Edgar Chambers; J. Scott Smith; Sara Gragg\nhttps://www.grains.k-state.edu/about/people/faculty/index.html"),
        ("University of Nebraska-Lincoln", "Food Science and Technology", "Dec. 1, 2026", "Andrew Benson; Devin Rose; Heather Hallen-Adams; Jeyam Subbiah\nhttps://foodscience.unl.edu/people/faculty/"),
        ("Oregon State University", "Food Science and Technology", "Dec. 1, 2026", "Yanyun Zhao; Jisun Park; Wytze van der Meer; Byron C. Brehm-Stecher\nhttps://foodsci.oregonstate.edu/foodsci/research/faculty-and-research-advisors"),
        ("Washington State University", "Food Science", "Jan. 10, 2027", "Carolyn Ross; Meijun Zhu; Charles Edwards; Girish Ganjyal\nhttps://sfs.wsu.edu/personnel/faculty-expertise-research-focus/"),
        ("Clemson University", "Food, Nutrition, and Packaging Sciences", "Jan. 15, 2027", "Feng Chen; Xiangyang Zhu; Paul Dawson; Angela Fraser\nhttps://www.clemson.edu/cafls/food-nutrition-packaging-sciences/directory/faculty.html"),
        ("University of Maryland, College Park", "Nutritional Sciences", "Dec. 1, 2026", "J. Glenn Morris; Kevin Fritsche; Prabhu S. B. Ramachandran; Chris D. H.\nhttps://agnr.umd.edu/about/directory/nutrition-food-science"),
        ("Tufts University", "Nutrition", "Dec. 1, 2026", "Friedman School faculty directory\nhttps://nutrition.tufts.edu/academics/faculty"),
        ("University of North Carolina at Chapel Hill", "Nutrition", "Dec. 1, 2026", "Naser Gharaibeh; Marcie H. S. Van; Ammer K. M.\nhttps://sph.unc.edu/nutr/unc-nutrition/nutr-our-faculty-and-staff/"),
        ("Auburn University", "Poultry Science and Food Safety", "Dec. 1, 2026", "Dianna Bourassa; Amit Morey; Ken Macklin; Jason O. Lee\nhttps://agriculture.auburn.edu/directory/index.php"),
        ("University of Arkansas", "Food Science", "Dec. 1, 2026", "Kristen Gibson; J. A. Marcy; Luke Howard\nhttps://food-science.uark.edu/people/"),
        ("Louisiana State University", "Food Science", "Dec. 1, 2026", "Marlene Janes; Witoon Prinyawiwatkul; Zhimin Xu\nhttps://www.lsu.edu/agriculture/nfs/people/index.php"),
        ("University of Tennessee, Knoxville", "Food Science and Technology", "Dec. 1, 2026", "Federico Harte; J. Howard; Robert M. Kerr\nhttps://foodscience.tennessee.edu/faculty-and-staff/"),
        ("University of Missouri", "Food, Nutrition and Exercise Sciences", "Dec. 1, 2026", "Andrew Clarke; Bongkosh Vardhanabhuti; Azlin Mustapha\nhttps://cafnr.missouri.edu/directory/?division=food-nutrition-and-exercise-sciences&role=faculty&research_area="),
        ("University of Connecticut", "Nutritional Sciences", "Dec. 15, 2026", "Ji-Young Lee; Dennis D'Amico; Nancy Rodriguez\nhttps://nusc.uconn.edu/faculty-2/"),
        ("University of Delaware", "Animal and Food Sciences", "Jan. 15, 2027", "Kali Kniel; Catherine Brown; Harshad Bhunia\nhttps://www.udel.edu/academics/colleges/canr/departments/animal-and-food-sciences/faculty-staff/"),
        ("University of Kentucky", "Dietetics and Human Nutrition", "Jan. 15, 2027", "James House; Kelly Webber; Lisa Cassis\nhttps://dhn.mgcafe.uky.edu/about-us/directory"),
        ("Texas Tech University", "Nutritional Sciences", "Dec. 1, 2026", "Naima Moustaid-Moussa; Min Du; Leslie Shen\nhttps://www.depts.ttu.edu/hs/ns/faculty.php"),
        ("University of Arizona", "Nutritional Sciences", "Dec. 1, 2026", "Cynthia Thomson; Carol Johnston; Federico I.\nhttps://snsw.arizona.edu/people/faculty"),
        ("University of Alabama at Birmingham", "Nutrition Sciences", "Dec. 1, 2026", "Barbara Gower; W. Timothy Garvey; SoJung Lee\nhttps://www.uab.edu/shp/nutrition/people"),
        ("Johns Hopkins University", "Human Nutrition", "Dec. 1, 2026", "Keith P. West Jr.; Parul Christian; Jessica Fanzo\nhttps://publichealth.jhu.edu/departments/international-health/people/faculty/primary-faculty-within-the-human-nutrition-program-area"),
        ("Harvard University", "Nutrition", "Dec. 1, 2026", "Frank Hu; Qi Sun; Walter Willett\nhttps://hsph.harvard.edu/department/nutrition/faculty-and-researcher-profiles/"),
        ("University of Washington", "Food Systems, Nutrition, and Health", "Dec. 1, 2026", "Adam Drewnowski; Marian Neuhouser; Anju Aggarwal\nhttps://foodsystems.uw.edu/about/nsp-faculty/"),
        ("University of California, Berkeley", "Metabolic Biology and Nutrition", "Dec. 1, 2026", "Kristine Madsen; faculty directory\nhttps://mbn.berkeley.edu/people/faculty"),
        ("University of Colorado Anschutz", "Nutrition", "Dec. 1, 2026", "Nancy F. Krebs; faculty directory\nhttps://medschool.cuanschutz.edu/pediatrics/sections/nutrition"),
        ("Drexel University", "Nutrition Sciences", "Jan. 15, 2027", "Stella Volpe; Mariana Chilton; faculty directory\nhttps://drexel.edu/cnhp/faculty/nutrition/"),
        ("University of Miami", "Public Health Sciences - Nutrition", "Dec. 1, 2026", "Sheah Rarback; faculty directory\nhttps://www.publichealth.med.miami.edu/faculty-and-staff/primary-faculty/index.html"),
        ("Oklahoma State University", "Nutritional Sciences", "Dec. 1, 2026", "Barbara Stoecker; faculty directory\nhttps://experts.okstate.edu/"),
        ("Colorado State University", "Food Science and Human Nutrition", "Dec. 1, 2026", "H. J. Swanson; faculty directory\nhttps://www.chhs.colostate.edu/fshn/about-us/faculty-and-staff/"),
        ("Mississippi State University", "Biochemistry, Nutrition and Health Promotion", "Dec. 1, 2026", "Juan L. Silva; faculty directory\nhttps://www.bchnhp.msstate.edu/people/directory/index.php"),
        ("South Dakota State University", "Dairy and Food Science", "Jan. 15, 2027", "Dairy and Food Science faculty directory\nhttps://www.sdstate.edu/dairy-food-science/our-people-dairy-food-science"),
        ("North Dakota State University", "Animal Sciences and Food Science", "Jan. 15, 2027", "Animal Sciences faculty directory\nhttps://www.ndsu.edu/agriculture/animal-sciences/animal-sciences-team"),
        ("University of Idaho", "Animal, Veterinary and Food Sciences", "Jan. 15, 2027", "College faculty directory\nhttps://www.uidaho.edu/agricultural-life-sciences/our-people"),
    ]
    add_table(
        document,
        programs,
        ("University Name", "Department Name", "PhD Application Deadline", "Prospective Faculty / Official Faculty Directory"),
        (Inches(1.65), Inches(2.1), Inches(1.0), Inches(4.4)),
    )
    document.add_paragraph().add_run("Funding note: Contact prospective faculty before applying to ask whether they expect to recruit and financially support doctoral students for Fall 2027.").italic = True
    document.save(OUTPUT_FILE)


if __name__ == "__main__":
    build_document()